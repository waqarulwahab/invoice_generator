
import streamlit as st
st.sidebar.page_link('pages/0_generate_invoice_DD.py',     label="Generate Invoice",               icon="🏡")
st.sidebar.page_link('pages/1_list_of_clients_projects.py',label="List of Clients / Projects List",icon="📓")    
st.sidebar.page_link('pages/2_add_new_client_project.py',  label="Add New Client/Project record",  icon="✒️")
st.sidebar.page_link('pages/3_doc_to_pdf.py',              label="Convert To PDF",                 icon="🖨️")

import streamlit as st
import docx
import os
import tempfile
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import base64
import weasyprint



import streamlit as st
import docx
import os
import tempfile
import pdfkit
from docx.shared import Inches
from PIL import Image

# Function to extract images from docx
def extract_images_from_docx(docx_filename):
    doc = docx.Document(docx_filename)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            img = Image.open(BytesIO(img_data))
            images.append(img)
    return images

# Function to convert docx to pdf
def convert_docx_to_pdf(docx_filename, images):
    doc = docx.Document(docx_filename)

    # Create a temporary HTML file
    html_content = ""
    for paragraph in doc.paragraphs:
        html_content += f'<p>{paragraph.text}</p>'
    for img in images:
        img_filename = tempfile.NamedTemporaryFile(delete=False, suffix='.png').name
        img.save(img_filename)
        img_tag = f'<img src="{img_filename}" width="{img.width}" height="{img.height}">'
        html_content += img_tag

    # Convert HTML to PDF using pdfkit with wkhtmltopdf backend
    pdf_filename = os.path.splitext(docx_filename)[0] + '.pdf'
    pdfkit.from_string(html_content, pdf_filename, options={'quiet': ''})

    # Clean up temporary files
    for img in images:
        img.close()

    return pdf_filename


st.title('DOCX to PDF Converter')

uploaded_file = st.file_uploader("Choose a DOCX file", type="docx")

if uploaded_file is not None:
    with open(uploaded_file.name, 'wb') as f:
        f.write(uploaded_file.getbuffer())

    st.markdown(f"Uploaded {uploaded_file.name}")

    # Extract images from DOCX
    images = extract_images_from_docx(uploaded_file.name)

    # Convert DOCX to PDF
    pdf_file = convert_docx_to_pdf(uploaded_file.name, images)

    # Display PDF download link
    with open(pdf_file, "rb") as f:
        pdf_bytes = f.read()
        st.download_button(
            label="Download PDF",
            data=pdf_bytes,
            file_name=os.path.basename(pdf_file),
            mime="application/pdf",
        )





